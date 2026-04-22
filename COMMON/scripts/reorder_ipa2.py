from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document(r"C:\Users\Admin\Desktop\New folder (2)\IPA 6.docx")

full_text = ""
for para in doc.paragraphs:
    full_text += para.text + "\n"

new_doc = Document()

kop = new_doc.add_paragraph()
kop.add_run("Soal Ujian Sekolah Ilmu Pengetahuan Alam\tT.P. 2025/2026")
kop.alignment = WD_ALIGN_PARAGRAPH.CENTER

new_doc.add_paragraph("Pilihan Berganda (50%)")
new_doc.add_paragraph("")

lines = full_text.split('\n')
current_q = ""
options = {}
output_paras = []

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    m = re.match(r'^(\d+)\.\s*(.*)', line)
    if m:
        if current_q and options:
            output_paras.append((current_q, dict(options)))
        num = m.group(1)
        current_q = f"{num}. {m.group(2)}"
        options = {}
    else:
        m = re.match(r'^([a-d])\.\s*(.*)', line, re.I)
        if m:
            options[m.group(1).lower()] = m.group(2)

if current_q and options:
    output_paras.append((current_q, dict(options)))

for q_text, opts in output_paras:
    p = new_doc.add_paragraph()
    p.add_run(q_text)
    
    for opt in ['a', 'b', 'c', 'd']:
        if opt in opts:
            p = new_doc.add_paragraph()
            p.add_run(f"{opt}. {opts[opt]}")

new_doc.save(r"C:\Users\Admin\Desktop\New folder (2)\IPA_Reorder.docx")
print("DONE! File saved: IPA_Reorder.docx")