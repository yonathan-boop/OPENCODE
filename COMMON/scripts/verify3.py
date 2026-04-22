from docx import Document

doc = Document(r"C:\Users\Admin\Desktop\New folder (2)\IPA_Reorder.docx")

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f"{text}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")